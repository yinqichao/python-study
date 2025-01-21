import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from itertools import groupby

xuqiu_excel = pd.read_excel('/Users/yinqichao/Downloads/新建村需求变更最终版（软件部分）.xlsx')

document = Document()

content_list = []

for i in range(171):
    changjing = xuqiu_excel.values[i, 0]
    zgn = xuqiu_excel.values[i, 1]
    miaoshu = xuqiu_excel.values[i, 2] + '：' + xuqiu_excel.values[i, 3]

    content_list.append({
        'cj': changjing,
        'zgn': zgn,
        'ms': miaoshu
    })

    # print(f'{i + 1}:{changjing},{zgn},{miaoshu}')
    #
    # table = document.add_table(rows=6, cols=2)
    # table.cell(0, 0).text = '功能名称'
    # table.cell(0, 1).text = changjing
    #
    # table.cell(1, 0).text = '子功能'
    # table.cell(1, 1).text = zgn
    #
    # table.cell(3, 0).text = '业务规格'
    # table.cell(3, 1).text = miaoshu

content_group = groupby(content_list, key=lambda x: (x['cj'], x['zgn']))

for key, group in content_group:
    print(key, group)

    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = '功能名称'
    table.cell(0, 1).text = key[0]

    table.cell(1, 0).text = '子功能'
    table.cell(1, 1).text = key[1]

    ms = ''

    for item in group:
        ms += item['ms']
        ms += '\n'

    table.cell(2, 0).text = '业务规格'
    table.cell(2, 1).text = ms

document.save('/Users/yinqichao/Downloads/result.docx')
