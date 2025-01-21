import pandas as pd
from docx import Document
from docx.oxml.ns import qn

cujinhui_excel = pd.read_excel('/Users/yinqichao/Downloads/促进会名单.xlsx')

document = Document()

document.styles['Normal'].font.name = u'仿宋_GB2312'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

for i in range(47):
    company = cujinhui_excel.values[i, 0]
    address = cujinhui_excel.values[i, 3]
    contact = cujinhui_excel.values[i, 6]
    phone = cujinhui_excel.values[i, 8]
    legal = cujinhui_excel.values[i, 1]
    print(f'{i + 1}:{company},{address},{contact},{phone},{legal}')

    table = document.add_table(rows=5, cols=4)
    # table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = '单位\n全称'
    table.cell(0, 1).merge(table.cell(0, 2))
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 1).text = company

    # table.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, 0).text = '单位\n地址'
    table.cell(1, 1).merge(table.cell(1, 2))
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(1, 1).text = address

    # table.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.cell(2, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 0).text = '联系人'
    table.cell(2, 1).text = contact

    # table.cell(3, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.cell(3, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(3, 0).text = '联系\n电话'
    table.cell(3, 1).text = phone

    table.cell(2, 2).merge(table.cell(3, 2))
    # table.cell(2, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 2).text = '法   定\n代表人'
    table.cell(2, 3).merge(table.cell(3, 3))
    table.cell(2, 3).text = legal

document.save('/Users/yinqichao/Downloads/会员单位列表.docx')
